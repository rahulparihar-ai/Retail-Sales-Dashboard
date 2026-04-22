"""Final Word Report — BEE Lab Group 4 — uses retail_sales_final.csv stats"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DB = RGBColor(0x1F,0x4E,0x79); MB = RGBColor(0x2E,0x75,0xB6)
WT = RGBColor(0xFF,0xFF,0xFF); DG = RGBColor(0x33,0x33,0x33)

def shd(cell, hex_c):
    tc=cell._tc; p=tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex_c); p.append(s)

def hd(doc, text, lvl=1):
    p=doc.add_heading(text, level=lvl)
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb=DB if lvl==1 else MB; r.font.name='Calibri'
        r.font.size=Pt(15 if lvl==1 else 12)

def para(doc, text, sz=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(sz); r.font.name='Calibri'; r.font.color.rgb=DG; return p

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.name='Calibri'; r.font.color.rgb=DG

def tbl_hdr(tbl, cols, bg="1F4E79"):
    row=tbl.rows[0]
    for i,col in enumerate(cols):
        c=row.cells[i]; c.text=col
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        shd(c,bg)
        r=c.paragraphs[0].runs[0]; r.font.bold=True; r.font.size=Pt(10)
        r.font.color.rgb=WT; r.font.name='Calibri'

def tbl_row(tbl, idx, vals, alt=False):
    row=tbl.rows[idx]; bg="EBF3FB" if alt else "FFFFFF"
    for i,v in enumerate(vals):
        c=row.cells[i]; c.text=str(v)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        shd(c,bg)
        r=c.paragraphs[0].runs[0]; r.font.size=Pt(10); r.font.name='Calibri'; r.font.color.rgb=DG

doc=Document()
for sec in doc.sections:
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
para(doc,"",sz=3)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("BEE LAB PROJECT REPORT"); r.font.size=Pt(28); r.font.bold=True
r.font.color.rgb=DB; r.font.name='Calibri'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Retail Sales Analytics: Excel & Power BI Dashboards")
r.font.size=Pt(14); r.font.color.rgb=MB; r.font.name='Calibri'
para(doc,"",sz=6)
ct=doc.add_table(rows=8,cols=2); ct.alignment=WD_TABLE_ALIGNMENT.CENTER; ct.style='Table Grid'
cover=[("Subject","Business & Economic Environment (BEE) Lab"),
       ("Group","Group 4"),("Year/Sem","B.Tech First Year — Semester II"),
       ("Date","April 2026"),("Member 1","[Name]  —  [Roll No.]"),
       ("Member 2","[Name]  —  [Roll No.]"),("Member 3","[Name]  —  [Roll No.]"),
       ("Faculty","[Professor Name]")]
for i,(lbl,val) in enumerate(cover):
    c0,c1=ct.rows[i].cells[0],ct.rows[i].cells[1]
    c0.text=lbl; c1.text=val; shd(c0,"1F4E79")
    shd(c1,"EBF3FB" if i%2==0 else "FFFFFF")
    r0=c0.paragraphs[0].runs[0]; r0.font.bold=True; r0.font.color.rgb=WT; r0.font.name='Calibri'; r0.font.size=Pt(10.5)
    r1=c1.paragraphs[0].runs[0]; r1.font.color.rgb=DG; r1.font.name='Calibri'; r1.font.size=Pt(10.5)
doc.add_page_break()

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────
hd(doc,"1. Introduction & Background")
para(doc,"Retail sales data contains rich insights about customer behavior, product performance, and regional trends. This project uses Microsoft Excel and Microsoft Power BI to transform a 300-order, 3-year (2020–2022) retail dataset into two interactive Business Intelligence dashboards that help stakeholders make fast, data-driven decisions.")

# ── 2. PROBLEM STATEMENT ──────────────────────────────────────────────────────
hd(doc,"2. Problem Statement")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
r=p.add_run('"Retail businesses generate large volumes of transactional data but often lack the tools to extract actionable insights. This project designs two interactive dashboards — one in Microsoft Excel and one in Microsoft Power BI — enabling stakeholders to monitor KPIs, compare category and regional performance, analyze profit margins, and forecast future sales without programming knowledge."')
r.italic=True; r.font.size=Pt(11); r.font.name='Calibri'; r.font.color.rgb=DG

# ── 3. OBJECTIVES ─────────────────────────────────────────────────────────────
hd(doc,"3. Objectives")
for obj in ["Build an interactive Excel Dashboard with Pivot Tables, PivotCharts, and Slicers for store sales performance monitoring.",
            "Develop a Power BI Dashboard with 16 DAX-powered KPI measures and a 3-month sales forecast.",
            "Identify top-performing and loss-making product sub-categories across all years.",
            "Perform Year-over-Year (YoY) and Month-over-Month (MoM) trend analysis.",
            "Deliver a professional project report and presentation for submission."]:
    bullet(doc,obj)

# ── 4. METHODOLOGY ────────────────────────────────────────────────────────────
hd(doc,"4. Methodology")
mt=doc.add_table(rows=6,cols=3); mt.style='Table Grid'
tbl_hdr(mt,["Phase","Tool","Activity"])
mdata=[("1 — Data Collection","Kaggle / CSV","Downloaded retail sales dataset (300 rows, 2020-2022)"),
       ("2 — Data Cleaning","Power Query","Fixed data types, removed nulls, standardized dates"),
       ("3 — Data Modeling","Power BI DAX","Created DateTable, Star Schema, 16 DAX measures"),
       ("4 — Dashboard Design","Excel + Power BI","Pivot Tables, Charts, Slicers, Visuals, Forecast"),
       ("5 — Analysis","Both Tools","KPI tracking, YoY comparison, trend analysis")]
for i,row in enumerate(mdata,1): tbl_row(mt,i,row,alt=(i%2==0))

# ── 5. DATASET ────────────────────────────────────────────────────────────────
hd(doc,"5. Dataset Description")
para(doc,"File: retail_sales_final.csv  |  Source: Kaggle Superstore (custom curated)",sz=11)
for s in ["Total Records: 300 orders across 3 years (2020, 2021, 2022)",
          "Categories: Technology ($250,980), Furniture ($74,164), Office Supplies ($9,251)",
          "Regions: West, South, East, Central",
          "Segments: Consumer, Corporate, Home Office",
          "21 loss-making orders (high-discount Furniture items)"]:
    bullet(doc,s)

# ── 6. DATA DICTIONARY ────────────────────────────────────────────────────────
hd(doc,"6. Data Dictionary")
dt=doc.add_table(rows=15,cols=3); dt.style='Table Grid'
tbl_hdr(dt,["Column","Type","Description"])
dd=[("Order ID","Text","Unique order identifier (CA-YYYY-XXXXXX)"),
    ("Order Date","Date","Date order placed (DD-MM-YYYY)"),
    ("Ship Date","Date","Date order shipped (DD-MM-YYYY)"),
    ("Customer Name","Text","Full name of customer"),
    ("Segment","Text","Consumer / Corporate / Home Office"),
    ("Region","Text","East / West / Central / South"),
    ("State","Text","US delivery state"),
    ("Category","Text","Technology / Furniture / Office Supplies"),
    ("Sub-Category","Text","Phones, Laptops, Chairs, Tables, Binders, etc."),
    ("Product Name","Text","Full product name"),
    ("Sales","Decimal","Gross revenue ($) — can be multi-unit"),
    ("Quantity","Integer","Units ordered"),
    ("Discount","Decimal","Discount rate (0.10 = 10%)"),
    ("Profit","Decimal","Net profit; negative = loss")]
for i,row in enumerate(dd,1): tbl_row(dt,i,row,alt=(i%2==0))

# ── 7. EXCEL DASHBOARD ────────────────────────────────────────────────────────
hd(doc,"7. Excel Dashboard — Store Sales Performance")
para(doc,"File: Store_Sales_Dashboard_FINAL.xlsx  |  6 sheets with pre-built charts and KPI cards",sz=11)
hd(doc,"7.1 Dashboard KPIs",lvl=2)
kd=[("Total Sales (2020-2022)","$334,395.01"),("Total Profit","$51,612.01"),
    ("Profit Margin","15.4%"),("Total Orders","300"),("Total Units Sold","699"),
    ("2020 Sales","$81,018.48"),("2021 Sales","$131,145.77  (+61.9% YoY)"),
    ("2022 Sales","$122,230.76  (-6.8% YoY)")]
kt=doc.add_table(rows=len(kd)+1,cols=2); kt.style='Table Grid'
tbl_hdr(kt,["KPI","Value"])
for i,row in enumerate(kd,1): tbl_row(kt,i,row,alt=(i%2==0))
hd(doc,"7.2 Sheets & Charts",lvl=2)
st=doc.add_table(rows=7,cols=2); st.style='Table Grid'
tbl_hdr(st,["Sheet","Contents"])
sd=[("DASHBOARD","KPI cards (2 rows × 5), 3 charts, segment & year tables"),
    ("SalesData","300-row raw data as formatted Excel Table"),
    ("PT_Category","Sales & Profit by Category × Region + Clustered Bar Chart"),
    ("PT_YearlyComparison","2020 vs 2021 vs 2022 monthly + Line Chart"),
    ("PT_TopProducts","Top 10 products by profit + Horizontal Bar Chart"),
    ("PT_SubCategory","All sub-categories ranked by profit (loss rows in red)")]
for i,row in enumerate(sd,1): tbl_row(st,i,row,alt=(i%2==0))
hd(doc,"7.3 Slicers (add manually in Excel)",lvl=2)
for s in ["Region Slicer — filters all charts simultaneously",
          "Category Slicer — isolates one product category",
          "Segment Slicer — Consumer / Corporate / Home Office",
          "Year Slicer — compare any single year or multi-year",
          "Date Timeline Slicer — drag to select custom date range"]:
    bullet(doc,s)

# ── 8. POWER BI DASHBOARD ─────────────────────────────────────────────────────
hd(doc,"8. Power BI Dashboard — Retail Sales & Forecast BI")
hd(doc,"8.1 Data Model (Star Schema)",lvl=2)
para(doc,"The data model uses two tables with a One-to-Many relationship:",sz=11)
for s in ["SalesData (Fact Table) — 300 rows, 14 columns",
          "DateTable (Dimension) — created via DAX CALENDAR() function",
          "Relationship: DateTable[Date]  ──1:*──  SalesData[Order Date]",
          "DateTable marked as official Date Table for time intelligence"]:
    bullet(doc,s)
hd(doc,"8.2 DAX Measures (16 total)",lvl=2)
dxd=[("Total Sales","SUM(SalesData[Sales])","Total revenue"),
     ("Profit Margin %","DIVIDE([Profit],[Sales],0)","% revenue kept as profit"),
     ("YTD Sales","TOTALYTD([Total Sales],DateTable[Date])","Cumulative Jan-to-date"),
     ("YTD Profit","TOTALYTD([Total Profit],DateTable[Date])","Cumulative profit YTD"),
     ("Sales Previous Year","CALCULATE([Sales],SAMEPERIODLASTYEAR(...))","Prior year benchmark"),
     ("YoY Sales Growth %","DIVIDE([Sales]-[Sales PY],[Sales PY],0)","Year-over-year % change"),
     ("MoM Sales Growth %","DIVIDE([Sales]-[Sales PM],[Sales PM],0)","Month-on-month % change"),
     ("Loss Orders Count","CALCULATE([Orders],Profit<0)","Count of unprofitable orders")]
dxt=doc.add_table(rows=len(dxd)+1,cols=3); dxt.style='Table Grid'
tbl_hdr(dxt,["Measure","Formula","Purpose"])
for i,row in enumerate(dxd,1): tbl_row(dxt,i,row,alt=(i%2==0))
hd(doc,"8.3 Report Pages",lvl=2)
pgd=[("Overview","5 KPI Cards, Bar Chart, Line Chart, Donut Chart, Map Visual"),
     ("Category Analysis","Matrix Table, Stacked Bar, Scatter Plot, KPI Card"),
     ("Regional View","Filled Map, Column Chart, Summary Table, Treemap"),
     ("Forecast","Line Chart + 3-month forecast (95% CI), MTD Sales, MoM Growth")]
pgt=doc.add_table(rows=len(pgd)+1,cols=2); pgt.style='Table Grid'
tbl_hdr(pgt,["Page","Visuals"])
for i,row in enumerate(pgd,1): tbl_row(pgt,i,row,alt=(i%2==0))
hd(doc,"8.4 Forecasting Setup",lvl=2)
for s in ["Visual: Line Chart with DateTable[Date] (Month level) on X-axis, Total Sales on Y-axis",
          "Analytics Pane → Forecast → Add",
          "Forecast length: 3 months  |  Confidence Interval: 95%  |  Seasonality: Auto",
          "Output: Dashed forecast line with shaded confidence band"]:
    bullet(doc,s)

# ── 9. KEY FINDINGS ───────────────────────────────────────────────────────────
hd(doc,"9. Key Findings & Business Insights")
for f in ["Technology dominates sales at $250,980 (75% of total) — driven by Laptops and Phones.",
          "2021 saw the highest YoY growth at +61.9%, followed by a -6.8% dip in 2022.",
          "West region leads in sales ($101,195), followed by South ($84,679).",
          "Laptops ($19,678) and Phones ($19,621) are the most profitable sub-categories.",
          "Bookcases and Tables show losses due to discounts exceeding 15–20%.",
          "Consumer segment contributes 47.6% of total revenue ($159,089).",
          "Forecast: Sales projected to recover in early 2023 based on seasonal trend."]:
    bullet(doc,f)

# ── 10. LIMITATIONS ───────────────────────────────────────────────────────────
hd(doc,"10. Limitations")
for l in ["Dataset uses US geography; may not reflect Indian retail patterns.",
          "Power BI forecasting is exponential smoothing only — no external factor modeling.",
          "300 rows is sufficient for learning; real dashboards use millions of records.",
          "No live database connection — CSV is static."]:
    bullet(doc,l)

# ── 11. CONCLUSION ────────────────────────────────────────────────────────────
hd(doc,"11. Conclusion")
para(doc,"This project demonstrates a complete Business Intelligence workflow — from raw CSV data to interactive dashboards using Microsoft Excel and Power BI. The Excel dashboard provides Pivot Table-based KPI monitoring with Slicer interactivity, while the Power BI dashboard delivers advanced DAX analytics, a proper Star Schema data model, and 3-month sales forecasting. Both tools together address all project objectives and provide a strong foundation for enterprise-level data analytics.")

# ── 12. REFERENCES ────────────────────────────────────────────────────────────
hd(doc,"12. References")
for ref in ["[1] Microsoft DAX Reference — docs.microsoft.com/en-us/dax/",
            "[2] Kaggle Superstore Dataset — kaggle.com/datasets/vivek468/superstore-dataset-final",
            "[3] Power BI Forecasting — learn.microsoft.com/en-us/power-bi/visuals",
            "[4] Excel PivotTable Guide — support.microsoft.com/en-us/office",
            "[5] Power BI Desktop — powerbi.microsoft.com/desktop"]:
    bullet(doc,ref)

OUT = r"e:\BEE\Section4_Report\BEE_Lab_Report_v2.docx"
doc.save(OUT)
print(f"Word report saved -> {OUT}")
